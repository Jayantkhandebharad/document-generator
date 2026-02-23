import logging
import os

from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
_log = logging.getLogger(__name__)
from docx.shared import Inches

from utils.formatter import (
    clear_document_body,
    force_legal_run_format_document,
    render_document_by_region,
    remove_trailing_empty_and_noise,
)
from utils.style_extractor import (
    _paragraph_has_bottom_border,
    extract_document_blueprint,
    extract_styles,
    load_extracted_styles,
    save_document_blueprint,
    save_extracted_styles,
)
from utils.zone_parser import extract_all_captions, parse_caption_structured, parse_regions

# Summons-style page margins (generous, like formal legal documents)
DEFAULT_TOP_MARGIN_IN = 1.25
DEFAULT_BOTTOM_MARGIN_IN = 1.25
DEFAULT_LEFT_MARGIN_IN = 1.25
DEFAULT_RIGHT_MARGIN_IN = 1.25


def _apply_default_margins(doc):
    """Ensure every section has at least default wide margins (proper spacing from page edges)."""
    try:
        for section in doc.sections:
            section.top_margin = Inches(DEFAULT_TOP_MARGIN_IN)
            section.bottom_margin = Inches(DEFAULT_BOTTOM_MARGIN_IN)
            section.left_margin = Inches(DEFAULT_LEFT_MARGIN_IN)
            section.right_margin = Inches(DEFAULT_RIGHT_MARGIN_IN)
    except Exception:
        pass


def _project_dir():
    return os.path.dirname(os.path.abspath(__file__))


def get_document_preview_text(docx_path: str) -> str:
    """Build a plain-text preview of the formatted DOCX for display before download.
    Paragraphs with only a bottom border (section underlines) are emitted as [SECTION_UNDERLINE]."""
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text and _paragraph_has_bottom_border(para):
            lines.append("[SECTION_UNDERLINE]")
        else:
            lines.append(text if text else "")
    return "\n\n".join(lines).strip()


def extract_and_store_styles(template_file) -> dict:
    """Extract styles from the uploaded DOCX and save to JSON. Returns the style schema."""
    doc = Document(template_file)
    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=_project_dir())
    blueprint = extract_document_blueprint(doc)
    save_document_blueprint(blueprint, base_dir=_project_dir())
    return schema


def process_document(generated_text, template_file):
    """
    Region-aware layout: layout dictates where content goes (no slot-fill for structure).
    Step 1 — Deterministic region extraction: caption, summons_intro, body, wherefore, signature, verification, footer.
    Step 2 — Render region by region in fixed order. Caption = table only. Divider only in caption and footer; never in body.
    """
    _log.info("process_document: start")
    project_dir = _project_dir()
    _log.info("process_document: loading template (Document)")
    doc = Document(template_file)
    _apply_default_margins(doc)

    _log.info("process_document: extract_styles + save")
    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=project_dir)

    # Step 1 — Deterministic region extraction and caption instances (no LLM)
    _log.info("process_document: parse_regions")
    regions = parse_regions(generated_text or "")
    schema["regions"] = regions
    _log.info("process_document: extract_all_captions")
    caption_instances = extract_all_captions(generated_text or "")
    if not caption_instances and regions.get("caption"):
        # Fallback: single caption at top from zones
        structured = parse_caption_structured(regions["caption"])
        if structured:
            caption_instances = [{"type": "summons_caption", "text": regions["caption"], "structured": structured}]
    schema["caption_instances"] = caption_instances

    _log.info("process_document: clear_document_body")
    clear_document_body(doc)
    # Step 2 — Render by region; caption instances in controlled positions (summons/complaint/footer)
    _log.info("process_document: render_document_by_region")
    render_document_by_region(doc, regions, caption_instances, schema)
    _log.info("process_document: force_legal_run_format_document")
    force_legal_run_format_document(doc)
    _log.info("process_document: remove_trailing_empty_and_noise")
    remove_trailing_empty_and_noise(doc)

    output_path = os.path.join(project_dir, "output", "formatted_output.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    _log.info("process_document: doc.save")
    doc.save(output_path)
    _log.info("process_document: get_document_preview_text")
    preview_text = get_document_preview_text(output_path)
    _log.info("process_document: done")
    return output_path, preview_text